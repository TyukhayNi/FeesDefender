#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Captura el delta de edición borrador ↔ versión firmada de un escrito (F10).

Señal más rica de la mejora continua: cada reescritura manual del letrado se
convierte en dato. Convención del despacho: la versión definitiva se guarda como
``<NOMBRE>_FIRMADO.docx`` en la misma carpeta que el borrador.

El script extrae el texto de ambos ``.docx`` (párrafos + celdas de tabla, con
``python-docx``), calcula las diferencias a nivel de párrafo (``difflib``) y
escribe ``<store>/<skill>/<ref>_delta.md`` con las correcciones del letrado
clasificadas en **añadido / suprimido / reescrito**.

⚠️ El delta contiene texto del escrito → **material sensible**. Vive en el store
central del repo (``data/_skill_logs/``, gitignored), **nunca** se empaqueta en
un ``.skill`` ni toca ``90_Notas personales``. Es work-product interno sin
anonimizar (decisión del despacho, plan v3 §16).

Uso:
  python scripts/capturar_delta.py <skill> <ref> --borrador ruta/ESCRITO.docx
                                   [--firmado ruta/ESCRITO_FIRMADO.docx]
Si se omite --firmado, se deriva como ``<borrador_stem>_FIRMADO.docx``.
"""
from __future__ import annotations

import argparse
import difflib
import os
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:  # pragma: no cover
    print("[capturar_delta] requiere python-docx (pip install python-docx).", file=sys.stderr)
    raise

_REPO = Path(__file__).resolve().parents[1]


def _store_dir(skill: str) -> Path:
    env = os.environ.get("FEESDEFENDER_SKILL_LOGS")
    base = Path(env) if env else _REPO / "data" / "_skill_logs"
    return base / skill


def extraer_parrafos(path: Path) -> list[str]:
    """Texto del .docx: párrafos del cuerpo + celdas de tabla, no vacíos."""
    doc = Document(str(path))
    bloques = [p.text.strip() for p in doc.paragraphs]
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                bloques.append(celda.text.strip())
    return [b for b in bloques if b]


def calcular_delta(borrador: list[str], firmado: list[str]) -> dict[str, list]:
    """Clasifica diferencias por párrafo en añadido / suprimido / reescrito."""
    sm = difflib.SequenceMatcher(None, borrador, firmado, autojunk=False)
    res: dict[str, list] = {"anadido": [], "suprimido": [], "reescrito": []}
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "insert":
            res["anadido"].extend(firmado[j1:j2])
        elif tag == "delete":
            res["suprimido"].extend(borrador[i1:i2])
        elif tag == "replace":
            res["reescrito"].append((borrador[i1:i2], firmado[j1:j2]))
    return res


def _bloque(titulo: str, items: list[str]) -> str:
    if not items:
        return f"### {titulo} (0)\n\n_— ninguno —_\n"
    out = [f"### {titulo} ({len(items)})", ""]
    out += [f"- {t}" for t in items]
    return "\n".join(out) + "\n"


def render_delta_md(skill: str, ref: str, delta: dict, borrador: Path, firmado: Path) -> str:
    reescritos = delta["reescrito"]
    partes = [
        f"# Delta de edición — {skill} — {ref}",
        "",
        "> Correcciones del letrado (borrador → firmado). Material de expediente "
        "sin anonimizar; no compartir ni empaquetar.",
        "",
        f"- Borrador: `{borrador.name}`",
        f"- Firmado:  `{firmado.name}`",
        f"- Resumen: {len(delta['anadido'])} añadidos · {len(delta['suprimido'])} "
        f"suprimidos · {len(reescritos)} reescritos",
        "",
        _bloque("Añadido por el letrado", delta["anadido"]),
        _bloque("Suprimido por el letrado", delta["suprimido"]),
        "### Reescrito (" + str(len(reescritos)) + ")",
        "",
    ]
    if not reescritos:
        partes.append("_— ninguno —_\n")
    else:
        for antes, despues in reescritos:
            partes.append("**Borrador:**")
            partes += [f"> {t}" for t in antes]
            partes.append("")
            partes.append("**Firmado:**")
            partes += [f"> {t}" for t in despues]
            partes.append("")
    return "\n".join(partes).rstrip() + "\n"


def capturar(skill: str, ref: str, borrador: Path, firmado: Path | None) -> Path:
    if firmado is None:
        firmado = borrador.with_name(f"{borrador.stem}_FIRMADO{borrador.suffix}")
    if not borrador.exists():
        raise FileNotFoundError(f"Borrador no encontrado: {borrador}")
    if not firmado.exists():
        raise FileNotFoundError(
            f"Versión firmada no encontrada: {firmado}\n"
            f"Convención: <NOMBRE>_FIRMADO.docx en la misma carpeta."
        )
    delta = calcular_delta(extraer_parrafos(borrador), extraer_parrafos(firmado))
    out_dir = _store_dir(skill)
    out_dir.mkdir(parents=True, exist_ok=True)
    destino = out_dir / f"{ref}_delta.md"
    destino.write_text(render_delta_md(skill, ref, delta, borrador, firmado), encoding="utf-8")
    return destino


def main(argv: list[str]) -> int:
    p = argparse.ArgumentParser(description="Captura el delta borrador↔firmado de un escrito.")
    p.add_argument("skill")
    p.add_argument("ref")
    p.add_argument("--borrador", required=True)
    p.add_argument("--firmado", default=None)
    args = p.parse_args(argv)
    destino = capturar(
        args.skill, args.ref, Path(args.borrador),
        Path(args.firmado) if args.firmado else None,
    )
    print(f"[capturar_delta] delta -> {destino}")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
