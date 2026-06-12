# -*- coding: utf-8 -*-
"""Genera MINUTA_AP_[REF].docx desde un JSON de datos.

Formato de la plantilla del despacho (confirmado): Arial 12, interlineado 1,25,
márgenes A4 2,5/2,5/3,5/2 cm (izquierdo amplio para anotar a mano), cabeceras de
bloque sombreadas con borde inferior, "[ PARA LEER EN SALA ]" en 9 pt, subpuntos
numerados jerárquicos, dos tablas de hechos (cabecera gris 25 %, filas alternas 5 %).

Markup inline admitido en los textos del JSON:
  **negrita**   _cursiva_

Uso:
  python gen_minuta.py datos.json salida.docx

Esquema del JSON: ver references/formato_minuta.md
"""
import json
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Arial"
GRAY_HEADER = "BFBFBF"   # cabecera de tabla (gris 25 %)
GRAY_ALT = "F2F2F2"      # filas alternas (gris 5 %)
GRAY_BLOCK = "D9D9D9"    # encabezado de bloque sombreado
DARKGRAY = RGBColor(0x40, 0x40, 0x40)


def _shade(el, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    el.append(shd)


def _bottom_border(p):
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "12"); b.set(qn("w:space"), "2"); b.set(qn("w:color"), "808080")
    pbdr.append(b); ppr.append(pbdr)


def _add_runs(p, text):
    for tk in re.split(r"(\*\*.*?\*\*|_.*?_)", text):
        if not tk:
            continue
        if tk.startswith("**") and tk.endswith("**"):
            p.add_run(tk[2:-2]).bold = True
        elif tk.startswith("_") and tk.endswith("_"):
            p.add_run(tk[1:-1]).italic = True
        else:
            p.add_run(tk)


class Minuta:
    def __init__(self):
        self.doc = Document()
        n = self.doc.styles["Normal"]
        n.font.name = FONT; n.font.size = Pt(12)
        n.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        n.paragraph_format.line_spacing = 1.25; n.paragraph_format.space_after = Pt(4)
        s = self.doc.sections[0]
        s.page_height = Cm(29.7); s.page_width = Cm(21.0)
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5); s.left_margin = Cm(3.5); s.right_margin = Cm(2.0)
        self._page_number(s)

    def para(self, text="", *, size=12, italic=False, bold=False, align=None, left=None,
             space_before=0, space_after=4, color=None):
        p = self.doc.add_paragraph()
        if align:
            p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(space_before); pf.space_after = Pt(space_after); pf.line_spacing = 1.25
        if left is not None:
            pf.left_indent = Cm(left)
        if text:
            _add_runs(p, text)
        for r in p.runs:
            r.font.size = Pt(size)
            if italic:
                r.italic = True
            if bold:
                r.bold = True
            if color:
                r.font.color.rgb = color
        return p

    def block_header(self, num, title, art):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(12); pf.space_after = Pt(6); pf.line_spacing = 1.0
        _shade(p._p.get_or_add_pPr(), GRAY_BLOCK); _bottom_border(p)
        r = p.add_run(f"{num} · {title}"); r.bold = True; r.font.size = Pt(12); r.font.name = FONT
        if art:
            r2 = p.add_run(f"   ({art})"); r2.italic = True; r2.font.size = Pt(10)

    def lectura(self):
        self.para("[ PARA LEER EN SALA ]", size=9, italic=True, color=DARKGRAY)

    def _set_cell(self, cell, text, *, bold=False, size=11, fill=None):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.1; p.paragraph_format.space_after = Pt(2)
        _add_runs(p, text)
        for r in p.runs:
            r.font.name = FONT; r.font.size = Pt(size)
            if bold:
                r.bold = True
        if fill:
            _shade(cell._tc.get_or_add_tcPr(), fill)

    def table(self, headers, rows, widths_cm, *, size=11):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
        for i, h in enumerate(headers):
            self._set_cell(t.rows[0].cells[i], f"**{h}**", bold=True, size=size, fill=GRAY_HEADER)
        for idx, row in enumerate(rows):
            cells = t.add_row().cells
            fill = GRAY_ALT if idx % 2 == 1 else None
            for i, val in enumerate(row):
                self._set_cell(cells[i], str(val), size=size, fill=fill)
        for row in t.rows:
            for i, w in enumerate(widths_cm):
                row.cells[i].width = Cm(w)
        return t

    def _page_number(self, section):
        p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
        ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve"); ins.text = "PAGE"
        f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
        run._r.append(f1); run._r.append(ins); run._r.append(f2)
        run.font.name = FONT; run.font.size = Pt(12)

    def save(self, path):
        self.doc.save(path)


def build(data, out_path):
    m = Minuta()
    cab = data["cabecera"]
    t = m.doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("MINUTA — AUDIENCIA PREVIA"); r.bold = True; r.font.size = Pt(14); r.font.name = FONT
    for linea in cab["lineas"]:
        m.para(linea, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    m.para("", space_after=4)

    for bloque in data["bloques"]:
        m.block_header(bloque["num"], bloque["titulo"], bloque.get("articulo", ""))
        if bloque.get("leer"):
            m.lectura()
        for it in bloque.get("items", []):
            # it: {"texto": str, "nivel": 0|1|2, "sub": false} ; nivel→sangría
            left = {0: 0.5, 1: 1.0, 2: 1.5}.get(it.get("nivel", 0), 0.5)
            m.para(it["texto"], left=left, space_before=4 if it.get("sub") else 0)
        # tablas opcionales dentro del bloque (fijación de hechos)
        for tab in bloque.get("tablas", []):
            m.para(f"**{tab['titulo']}**", size=11, space_before=6, space_after=2)
            m.table(tab["cabecera"], tab["filas"], tab["anchos_cm"], size=tab.get("size", 11))
        if bloque.get("nota"):
            m.para(bloque["nota"], size=10, space_before=6, color=DARKGRAY)
    m.save(out_path)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Uso: python gen_minuta.py datos.json salida.docx", file=sys.stderr)
        sys.exit(2)
    with open(sys.argv[1], encoding="utf-8") as fh:
        datos = json.load(fh)
    print("GUARDADO:", build(datos, sys.argv[2]))
