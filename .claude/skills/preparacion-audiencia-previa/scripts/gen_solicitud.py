# -*- coding: utf-8 -*-
"""Genera SOLICITUD_PRUEBA_[REF].docx con el formato EXACTO de la plantilla del
despacho (escrito procesal; alineado con `escritos-judiciales`):

- Times New Roman 12 pt; márgenes 2,5 cm; A4.
- Tabla de referencia en cabecera (Mi ref. / Ref. procurador).
- Párrafos justificados, interlineado 1,5, espaciado anterior 6 pt.
- Nombres de personas en MAYÚSCULAS NEGRITA; DNI en negrita; palabras clave en
  negrita (DIGO, PROPOSICIÓN DE LAS PRUEBAS, DOCUMENTAL, TESTIGO, CITACIÓN JUDICIAL,
  SUPLICO, firma). "en calidad de" NO va en negrita.
- Número de página centrado (TNR 12) en el pie.

Uso:
  python gen_solicitud.py datos.json salida.docx
"""
import json
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TNR = "Times New Roman"


def _base(doc):
    n = doc.styles["Normal"]
    n.font.name = TNR
    n.font.size = Pt(12)
    n.element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    pf = n.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    s = doc.sections[0]
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.5)
    p = s.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    f1 = OxmlElement("w:fldChar")
    f1.set(qn("w:fldCharType"), "begin")
    ins = OxmlElement("w:instrText")
    ins.set(qn("xml:space"), "preserve")
    ins.text = "PAGE"
    f2 = OxmlElement("w:fldChar")
    f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1)
    run._r.append(ins)
    run._r.append(f2)
    run.font.name = TNR
    run.font.size = Pt(12)


def _p(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, left=None):
    par = doc.add_paragraph()
    par.alignment = align
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(0)
    if left is not None:
        par.paragraph_format.left_indent = Cm(left)
    return par


def _run(par, text, bold=False):
    r = par.add_run(text)
    r.font.name = TNR
    r.font.size = Pt(12)
    r.bold = bold
    return r


def _ref_table(doc, ref, ref_proc):
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    t._tbl.tblPr.append(borders)
    data = [("Mi ref:", ref), ("Ref. procurador:", ref_proc)]
    for i, (k, v) in enumerate(data):
        for j, val in enumerate((k, v)):
            cell = t.rows[i].cells[j]
            cell.text = ""
            rr = cell.paragraphs[0].add_run(val)
            rr.font.name = TNR
            rr.font.size = Pt(12)
            if j == 0:
                rr.bold = True
        t.rows[i].cells[0].width = Cm(3.3)
        t.rows[i].cells[1].width = Cm(13.0)
    return t


def build(data, out_path):
    doc = Document()
    _base(doc)
    _ref_table(doc, data.get("ref", ""), data.get("ref_procurador", ""))

    _p(doc, align=WD_ALIGN_PARAGRAPH.LEFT)
    pj = _p(doc, align=WD_ALIGN_PARAGRAPH.LEFT)
    _run(pj, data["juzgado"], bold=True)
    pp = _p(doc, align=WD_ALIGN_PARAGRAPH.LEFT)
    _run(pp, data["procedimiento"], bold=True)

    al = _p(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _run(al, "AL JUZGADO", bold=True)

    comp = _p(doc)
    _run(comp, data["procurador"], bold=True)
    _run(comp, ", Procurador de los Tribunales, actuando en nombre y representación de ")
    _run(comp, data["cliente"], bold=True)
    _run(comp, ", cuya representación consta acreditada en los autos, ante este Juzgado comparezco y como mejor en Derecho proceda, ")
    _run(comp, "DIGO,", bold=True)

    intro = _p(doc)
    _run(intro, "Que habiendo sido notificada a esta representación la providencia por la que se convoca a las partes al acto de audiencia previa conforme a lo previsto en el artículo 414.1 LECiv en correlación con el artículo 429 LECiv, mediante el presente escrito esta representación formula ")
    _run(intro, "PROPOSICIÓN DE LAS PRUEBAS", bold=True)
    _run(intro, " de las que intentará valerse, solicitando que sean admitidas en su totalidad por guardar relación con la tutela judicial pretendida y al ser pertinentes y útiles al guardar relación con el objeto del proceso.")

    for i, pr in enumerate(data["pruebas"], start=1):
        tipo = pr["tipo"]
        if tipo in ("documental", "mas_documental"):
            par = _p(doc, left=0.75)
            etiqueta = "DOCUMENTAL" if tipo == "documental" else pr.get("etiqueta", "MÁS DOCUMENTAL")
            _run(par, str(i) + ". ")
            _run(par, etiqueta + " ", bold=True)
            _run(par, pr["texto"])
        elif tipo == "testigo":
            par = _p(doc, left=0.75)
            _run(par, str(i) + ". ")
            _run(par, "DECLARACIÓN", bold=True)
            _run(par, " en calidad de ")
            _run(par, "TESTIGO", bold=True)
            _run(par, " de ")
            _run(par, pr["nombre"], bold=True)
            _run(par, ", con DNI ")
            _run(par, pr.get("dni", ""), bold=True)
            _run(par, ", " + pr.get("descripcion", "") + ".")
            cit = _p(doc, left=0.75)
            _run(cit, "Solicitamos su ")
            _run(cit, "CITACIÓN JUDICIAL", bold=True)
            _run(cit, " en " + pr.get("citacion", ""))
            if pr.get("movil"):
                _run(cit, ", teléfono móvil " + pr["movil"])
            if pr.get("email"):
                _run(cit, ", email " + pr["email"])
            _run(cit, ".")
        elif tipo == "interrogatorio":
            par = _p(doc, left=0.75)
            _run(par, str(i) + ". ")
            _run(par, "INTERROGATORIO", bold=True)
            _run(par, " " + pr["texto"])
        elif tipo == "oficio":
            par = _p(doc, left=0.75)
            _run(par, str(i) + ". ")
            _run(par, "OFICIO", bold=True)
            _run(par, " " + pr["texto"])
        else:
            par = _p(doc, left=0.75)
            _run(par, str(i) + ". " + pr.get("texto", ""))

    sup = _p(doc)
    _run(sup, "SUPLICO AL JUZGADO", bold=True)
    _run(sup, " tenga por presentado este escrito junto con las copias que lo acompañan y acuerde de conformidad con su contenido.")

    just = _p(doc)
    _run(just, "Es justicia lo que respetuosamente pido, a " + data.get("fecha", ""))

    firma = _p(doc, align=WD_ALIGN_PARAGRAPH.LEFT)
    _run(firma, "LTDO. " + data.get("letrado", "") + "                    PROC. " + data.get("procurador_firma", data.get("procurador", "")), bold=True)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Uso: python gen_solicitud.py datos.json salida.docx", file=sys.stderr)
        sys.exit(2)
    with open(sys.argv[1], encoding="utf-8") as fh:
        datos = json.load(fh)
    print("GUARDADO:", build(datos, sys.argv[2]))
